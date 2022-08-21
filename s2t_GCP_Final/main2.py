from docx import Document
import pyaudio
import wave
from tkinter import *
import threading
from datetime import datetime
from tkinter.filedialog import askopenfile
import os
from Merge import merge_code

GOOGLE_CLOUD_SPEECH_CREDENTIALS = r"""{
  "type": "service_account",
  "project_id": "amrproject-341506",
  "private_key_id": "ab0f62cbbcc323a6b013188ed3446e7eb002fdb7",
  "private_key": "-----BEGIN PRIVATE KEY-----\nMIIEvQIBADANBgkqhkiG9w0BAQEFAASCBKcwggSjAgEAAoIBAQDvEy9BU+eTefM4\ngo85aV8/+yMhwZaXjb6a1Ghh8eNjFZdJO1x890ighYnQLtuMgltqKbMrBjOECiDO\nscdsie3sfogCloRArEr+s07cR32YdWMPcV96CWQYWbWsxmgU+0SsdbXUD9qfhICK\nyi3JYitYdjefB0iMP2oyQl0LJaZiehtlaS+XwAACusubV/oVmcXsStQY0UudxnRR\nN6w4k9I9WAbm9hua0bX4j6dedHgqZySfPxBkNDOpsQj+7pN2fidQpAlRBCU+IBI4\n/a4M5DEQjl6l5ZPiCjpWbuwYXek8/87BbeBpcMjgoCVdEWP/GoHWSsUpPMPSadN0\nXAelPbVpAgMBAAECggEAM2ZCgJ8TaKeNXRbkyAkZi13o/bbrM4Yr9om4L2ab5o+b\nyQjUJ6HaoXhUWsfYN+tG0BpLBPL+XGL5Drc2EuXMm03Q3A2UPXtNXFkcmUtde0ux\nkn5jco9WT83pu02kzu3+o71i7EbEUAtiv3QcR/GtNAWmi7A7Lrvrjwapjl332a/Q\nknAgY6guAoIUjsq8FKFwicn22i/VsttutQvr+sCfaG3G6JSrmu8XNERstVujBuwM\n7zhLB2LPoY/qeMe+E1YCSUOF/qzeGoO0wDCJFcZdRdl7rQ4BuLVzE9Mq7c5xo5Vv\nsb/x67sHYuyuHmBA45iM1xmPet+5bhDow/UZCQTgqQKBgQD8NsJJRlZYmbYOt7qf\n6qrmnxi9sbyq6PyMo4AcorDrtr/ZIO0RePh6NV+A60HjtxMvo1OlLnmpA2HG+AR3\nSnVS+ayuDNoV2WQQOy6aYRVKZkzWadj5yaCV1pE7WGM+7DgXu5AGCHba0rStVWd9\nh0z5vVZFH7ezrmqiGBJ5gsqScwKBgQDyqe76mTS3PbPB/zU6fh59a+RZaCiBso0x\nQQTnzS568lnVyw69wd4i9JPz1dk32f/Dy7/q+sFayTO/43sxAqToRBxAvG1FXcwe\naNY2VO85ALCgIpUJT/z/EQf5Vc+5W+FGBu1Nd3PAFn6Lohe5/kNUuN7eXL9z0a99\nJTztjjq1swKBgQCMFxZfgPQhDm/zSNYeUUhQBPrHEtxKylLNA1v7i9pdcwqo3gMP\n+3N43gCIKSRWCbBce4nQbDojSL+cttI0OGVmi0wA+R6BmFheM6E1GsuqoZv4VJ6y\nT+4kjCPrJLMrliWp+QKMd1MdRIAa7x8muGpnJWWvy2hMI4sTMe/BEobWRQKBgCP7\nmtpCGZUKaztbEpOH3BGRDB1N2qPQzCr/jzCoNjo48aVIJPlyY/Qe/Elt7nXft5Nb\nBSxc3osNxsSvgLb1w2pFn0LiLfnZPMlNlmYcFhUHyykTt7HX0JYp+lq4pMLBoxf1\nMv8HtXxIWZ9u+8GyuF05xKKIckNEVpdiDDTVoFMNAoGAWF7Oa/pxv1e3N+xGBKwV\noU4cCeTlhJfxQK486D1LZwb10Y5pSZjKVFfbxW4/LHhLAXfpETriKIb4TGpeURpf\nhzBjSDuVNOF9EEiMAB33JKIlrsa996fm+RgQmaG2TiFC3tZOJ4j4uOug4WrfRjhg\nDDTM8BR9y4tlDuPFZzb7icU=\n-----END PRIVATE KEY-----\n",
  "client_email": "speechproject@amrproject-341506.iam.gserviceaccount.com",
  "client_id": "114544725460519108555",
  "auth_uri": "https://accounts.google.com/o/oauth2/auth",
  "token_uri": "https://oauth2.googleapis.com/token",
  "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
  "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/speechproject%40amrproject-341506.iam.gserviceaccount.com"
}
"""

sample_format = pyaudio.paInt16
channel = 1
fs = 44100
frames_per_buffers = 1024
audio = pyaudio.PyAudio()

os.makedirs('ConvertWord', exist_ok=True)
os.makedirs('TrackRecord', exist_ok=True)


def wav_to_text(name):
    print(name)
    os.makedirs('{}'.format(name.split('.')[1]), exist_ok=True)
    song = AudioSegment.from_wav(name)
    r = sr.Recognizer()
    r.energy_threshold = 100
    text = []
    chunks = split_on_silence(song, min_silence_len=500, silence_thresh=-80)
    for i, chunk in enumerate(chunks):
        filename_ = '{}/chunk_'.format(name.split('.')[1]) + str(i) + '.wav'
        # Store the sliced audio file to the defined path
        chunk.export(filename_, format="wav")
        with sr.AudioFile(filename_) as source:
            # r.adjust_for_ambient_noise(source)
            audio_ = r.record(source)
            try:
                # rec = r.recognize_google_cloud(audio_, language='vi-VN',
                #                                credentials_json=GOOGLE_CLOUD_SPEECH_CREDENTIALS)
                rec = r.recognize_google(audio_, language='vi-VN', show_all=True)
                text.append(rec)
                # If Google could not understand the audio
            except sr.UnknownValueError:
                print("Could not understand audio")
            # If the results cannot be requested from Google.
            # Probably an internet connection error.
    text
    return text



class App(Tk):
    def __init__(self):
        super().__init__()

        # configure the root windows
        self.title("Voice Recorder App")
        self.geometry("800x400+500+100")

        self.canvas = Canvas(self, height=400, width=800, bd=0, highlightthickness=0, relief="ridge")
        self.canvas.place(x=0, y=0)

        self.background_img = PhotoImage(file=f"image/background.png")
        self.round = self.canvas.create_image(400.0, 200.0, image=self.background_img)
        self.header = self.canvas.create_text(400.0, 80.0, text="STECH Voice Recorder",
                                              fill="#F9D205",
                                              font=("Montserrat-Bold", int(30.0))
                                              )

        # Button
        self.start_img = PhotoImage(file=f"image/start.png")
        self.start = Button(self, image=self.start_img, borderwidth=0, highlightthickness=0, relief='flat')
        self.start.place(x=118, y=200, width=172, height=58)

        self.stop_img = PhotoImage(file=f"image/stop.png")
        self.stop = Button(self, image=self.stop_img, borderwidth=0, highlightthickness=0, relief='flat')
        self.stop.place(x=318, y=200, width=172, height=58)

        self.convert_img = PhotoImage(file=f"image/convert.png")
        self.convert = Button(self, image=self.convert_img, borderwidth=0, highlightthickness=0, relief='flat')
        self.convert.place(x=518, y=200, width=172, height=58)

        self.start.bind("<Button-1>", self.startrecording)
        self.convert.bind("<Button-1>", self._convert)

        # When start app
        self.info = self.canvas.create_text(400.0, 342.5, text="Click Start to Record",
                                            fill="#F9D205",
                                            font=("Montserrat-Bold", int(16.0)))
        self.stop.config(state='disabled')

    def startrecording(self, event):
        self.isrecording = True
        self.start.config(state='disabled')
        self.stop.config(state='normal')
        self.convert.config(state='disabled')
        self.start.unbind("<Button-1>")
        self.stop.bind("<Button-1>", self.stoprecording)
        self.convert.unbind("<Button-1>")
        self.canvas.itemconfig(self.info, text="Recording...")
        t = threading.Thread(target=self._record)
        t.start()

    def stoprecording(self, event):
        self.isrecording = False
        self.start.config(state='normal')
        self.stop.config(state='disabled')
        self.convert.config(state='normal')
        self.start.bind("<Button-1>", self.startrecording)
        self.stop.unbind("<Button-1>")
        self.convert.bind("<Button-1>", self._convert)

    def _record(self):
        now = datetime.now()
        name = now.strftime("%d-%m-%Y %H %M %S")

        self.audio_file = "./TrackRecord/{}.wav".format(name)

        doc = Document()
        self.word_file = "./ConvertWord/{}.docx".format(name)

        stream = audio.open(format=sample_format,
                            channels=channel,
                            rate=fs,
                            input=True,
                            frames_per_buffer=frames_per_buffers)

        frames = []
        while self.isrecording:
            data = stream.read(frames_per_buffers)
            frames.append(data)

        sound_file = wave.open(self.audio_file, "wb")
        sound_file.setnchannels(channel)
        sound_file.setsampwidth(audio.get_sample_size(sample_format))
        sound_file.setframerate(fs)
        sound_file.writeframes(b''.join(frames))
        sound_file.close()

        text = merge_code(self.audio_file)
        doc.add_heading('Conversation', level=0)
        for sen in text:
            doc.add_paragraph(sen.strip())
        doc.save(self.word_file)
        self.canvas.itemconfig(self.info, text="Stopped Record. File Saved At Outputs Folder")

    def _convert(self, event):
        doc = Document()
        file = askopenfile(mode='r', filetypes=[("Audio File", ".wav")])

        if file is not None:
            file_name = file.name.split('/')[-1]
            word_file = os.path.splitext(file_name)
            word_name = "./ConvertWord/{}.docx".format(word_file[0])
            file_name = "./TrackRecord/" + file_name

            text = merge_code(file_name)

            doc.add_heading('Conversation', level=0)
            for sen in text:
                doc.add_paragraph(sen.strip())
            doc.save(word_name)
            self.canvas.itemconfig(self.info, text="Converted. File Saved At Outputs Folder")
        else:
            self.canvas.itemconfig(self.info, text="Could not understand")


if __name__ == "__main__":
    app = App()
    app.mainloop()
