# !pip install google-cloud-speech
# !pip install google-cloud-storage
# !pip install pydub
# !pip install speechbrain
# !pip install pyannote.audio==2.0.1
# !pip install pyannote.core
# !pip install underthesea
# !pip install docx-python
import os
from pydub import AudioSegment, silence
import wave
from google.cloud import speech
from google.cloud import storage
from concurrent.futures import ThreadPoolExecutor
import io

from pyannote.audio import Pipeline

from docx import Document
from underthesea import pos_tag as pot
import re

pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization")

os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = 'Speech2Text.json'

## Tạo thư mục cho từng tác vụ
Thu_muc_cho_audio_cat_theo_giong_noi = 'Splited_speaker'
Thu_muc_cho_audio_cat_theo_silence   = 'Splited_Silence'

os.makedirs(Thu_muc_cho_audio_cat_theo_giong_noi ,exist_ok= True)
os.makedirs(Thu_muc_cho_audio_cat_theo_silence   ,exist_ok= True)
os.makedirs('ConvertWord', exist_ok=True)

def Create_Word(Audio_Input,final_text):
    if '.wav' in Audio_Input:
        doc = Document()
        file_name = Audio_Input.split('/')[-1]
        word_file = os.path.splitext(file_name)
        word_name = "./ConvertWord/{}.docx".format(word_file[0])
    elif 'ConvertWord' not in Audio_Input:
        if '.docx' not in Audio_Input:
            word_name = Audio_Input + '.docx'
        word_name = "./ConvertWord/{}.docx".format(word_name)
    else:
        if '.docx' not in Audio_Input:
            word_name = Audio_Input + '.docx'
        word_name = Audio_Input
    doc.add_heading('Conversation',level=0)
    for sen in final_text:
        doc.add_paragraph(sen.strip())
    doc.save(word_name)

# Split Speaker
# def Split_speaker(audio_file,time_ok=False):
#     Audio = AudioSegment.from_file(audio_file)
#     name = Thu_muc_cho_audio_cat_theo_giong_noi + '/' + audio_file.split('/')[-1].replace('.wav', '')

#     os.makedirs(name, exist_ok=True)
#     hashDict = dict()
#     timeDict = dict()
#     diarization = pipeline(audio_file)
    
#     i,start,end,lab = 0,0,0,0
#     for segment,track,label in diarization.itertracks(yield_label = True):
#         if label != lab:
#             Audio_speaker_i = Audio[start * 1000 : end * 1000 + 299]
#             Audio_speaker_i.export(name + '/' + str(i) + '.wav', format = 'wav')
#             hashDict[name + '/' +str(i) + '.wav'] = lab
#             if time_ok==True:
#                 timeDict[(start * 1000 , end * 1000 + 299)] = lab
#             i += 1
#             start = segment.start
#             end = segment.end
#             lab = label
            
#         elif label == lab:
#             end = segment.end
#     # chạy vòng cuối
#     Audio_speaker_i = Audio[start * 1000 : end * 1000 + 299]
#     Audio_speaker_i.export(name +'/' + str(i) + '.wav', format = 'wav')
#     hashDict[name + '/' +str(i) + '.wav'] = lab
#     if time_ok==True:
#             timeDict[(start * 1000 , end * 1000 + 299)] = lab
#             timeDict.pop((0,0 + 299))
            
#     os.remove(name + '/' + str(0) + '.wav')
#     hashDict.pop(name + '/' +str(0) + '.wav')
#     if time_ok==True:
#         return timeDict,hashDict
#     return hashDict
def Split_speaker(path,time_ok = False):
    Audio = AudioSegment.from_file(path)
    name = Thu_muc_cho_audio_cat_theo_giong_noi + '/' + path.split('/')[-1].replace('.wav', '')
    start1 = []
    stop1 = []
    label1 = []
    diarization = pipeline(path)
    for segment,track,label in diarization.itertracks(yield_label = True):
      start1.append(segment.start)
      stop1.append(segment.end)
      label1.append(label)
    starts = [start1[0]]
    stops = []
    labels = [label1[0]]

    for i in range(1,len(label1)):
      if label1[i] == labels[-1]:
        pass
      else:
        starts.append(start1[i])
        labels.append(label1[i])
        stops.append(stop1[i-1])
    stops.append(stop1[-1])
    duration = []

    for i in range(len(labels)):
      duration.append(stops[i]-starts[i])
    start_ = [starts[0]]
    stop_ = []
    label_ = [labels[0]]
    for i in range(1,len(labels)):
      if duration[i] < 0.5:
        pass
      else:
        start_.append(starts[i])
        label_.append(labels[i])
        stop_.append(stops[i-1])
    stop_.append(stops[-1])
    hashDict = dict()
    TimeDict = dict()
    os.makedirs(name,exist_ok=True)
    for i in range(len(label_)):
        Audio_speaker_i = Audio[start_[i]*1000:stop_[i]*1000 + 299]
        Audio_speaker_i.export(name + '/' + str(i) + '.wav',format='wav')
        hashDict[name + '/' + str(i) + '.wav'] = label_[i]
        if time_ok == True:
            TimeDict[(start_[i]*1000,stop_[i]*1000 + 299)] = label_[i]
    if time_ok==True:
        return TimeDict,hashDict
    return hashDict

# Upload Audio to google Cloud, do this if audio > 1m
def upload_blob(bucket_name, source_file_name, destination_blob_name):
    """Uploads a file to the bucket."""
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(destination_blob_name)

    blob.upload_from_filename(source_file_name)

    print(
        f"File {source_file_name} uploaded to {destination_blob_name}."
    )

# Một số Function hỗ trợ xử lý văn bản đầu ra
def is_digit(word):
    try:
        int(word)
        return True
    except ValueError:
        pass
    return False

def ConvertDate(text):
    month=' tháng '
    year=' năm '
    for index in range(0,len(text)):
        try:
            if (text.index(month,index)==index):
                dateNum = text[index -1]
                monthNum = text[index + len(month)]
                if is_digit(dateNum) and is_digit(monthNum):
                    text=text[:index] + text[index+len(month)-1:]
                    temp = list(text)
                    temp[index]='/'
                    text = "".join(temp)
        except Exception as e:
            if str(e) in 'substring not found':
                pass
            else:
                raise e
        try:
            if (text.index(year,index)==index):
                monthNum = text[index -1]
                yearNum = text[index + len(year)]
                if is_digit(monthNum) and is_digit(yearNum):
                    text=text[:index] + text[index+len(year)-1:]
                    temp = list(text)
                    temp[index]='/'
                    text = "".join(temp)
        except Exception as e:
            if str(e) in 'substring not found':
                pass
            else:
                raise e
    return text

# Các Function xử dụng GCP API thực hiện speech to text
def frame_rate_channel(audio_file_name):
    print(audio_file_name)
    with wave.open(audio_file_name, "rb") as wave_file:
        frame_rate = wave_file.getframerate()
        channels = wave_file.getnchannels()
        return frame_rate,channels

# Audio (< 1m) to text 
def Transcribe_Short_Audio(Audio_wav,config_wav_enhanced):
    client = speech.SpeechClient()
    with io.open(Audio_wav, "rb") as audio_file:
        content = audio_file.read()

    audio = speech.RecognitionAudio(content=content)
    # print(type(audio))
    
    response = client.recognize(config=config_wav_enhanced, audio=audio)
    text = []
    for i, result in enumerate(response.results):
        alternative = result.alternatives[0]
        alter=ConvertDate(alternative.transcript )
        text.append(alter)
    return text

# Audio (> 1m) to text 
def Transcribe_Long_Audio(Audio_wav,config_wav_enhanced,
                        bucket_name = 'speech_to_text_stech',
                        Name = 'Audio_wav'):
    client = speech.SpeechClient()
    Audio_name = Audio_wav.split('/')[-1].split('.')[0]
    
    upload_blob(bucket_name,Audio_wav,Name)
    
    media_uri = "gs://{}/{}".format(bucket_name,Name)
    long_audi_wav = speech.RecognitionAudio(uri=media_uri)
    
    operations = client.long_running_recognize(
        config = config_wav_enhanced,
        audio = long_audi_wav
    )
    
    response = operations.result(timeout=90)

    text = []
    
    for i, result in enumerate(response.results):
            alternative = result.alternatives[0]
            alter=ConvertDate(alternative.transcript)
            text.append(alter)
    return text

# Sử dụng Google API để dịch từng Audio sang text
def func_sort(e):
    return e[0]

def GCP_s2t(alist):
    index=list(alist[0])[0]
    sample_file_config=list(alist[0][index])[0]
    
    rate,channel=frame_rate_channel(sample_file_config)

    config = speech.RecognitionConfig(
            sample_rate_hertz = rate,
            language_code = 'vi-VN',
            audio_channel_count=channel,
            enable_automatic_punctuation=True,
            )

    config_noPunc = speech.RecognitionConfig(
            sample_rate_hertz = rate,
            language_code = 'vi-VN',
            audio_channel_count=channel,
            )

    def thread_function(index,file_addr):
        # global thTranscribeConfig
        try:
            text=Transcribe_Short_Audio(file_addr,thTranscribeConfig)
        except:
            text = Transcribe_Long_Audio(file_addr,thTranscribeConfig)
        x = index,text
        return x

    result_dict = {'Speaker':[], 'co dau':[],'khong dau':[],'silence':[]}
    for adict in alist:
        for speaker,value in adict.items():
                file_addr = list(adict[speaker].keys())
                silent_time = list(adict[speaker].values())
                index = list(range(len(file_addr)))
                thTranscribeConfig = config
                with ThreadPoolExecutor(8) as executor:
                    futures = executor.map(thread_function,index,file_addr)
                    index_text = [x for x in futures]
                index_text.sort(key=func_sort)
                text = [text for _,text in index_text]

                thTranscribeConfig = config_noPunc
                with ThreadPoolExecutor(8) as executor:
                    futures2 = executor.map(thread_function,index,file_addr)
                    index_text2 = [y for y in futures2]
                index_text2.sort(key=func_sort)
                text2 = [text2 for _,text2 in index_text2]

                result_dict['Speaker'].extend([speaker] * len(silent_time))
                result_dict['co dau'].extend(text)
                result_dict['khong dau'].extend(text2)
                result_dict['silence'].extend(silent_time)
    return result_dict

# Function Kết hợp Text có dấu theo Google API và silence
def Handle_Text_Dict(adict,S_Punc_ok=False,XD=False,newLine=2):
    adict = list(adict.values())
    Spr,SL = 0,0
    ReDict = {'Speaker':[], 'co dau':[],'khong dau':[]}
    if XD == True:
        ReDict['Xuong Dong'] = []
    # Silent_Punc = []
    Flag = False
    # commitCount=0

    for Speaker,Punc,NoPunc,Silent in zip(*adict):
        #  Vì Punc,NoPunc là list
        new_Punc = str()
        for i in Punc:
            new_Punc += i
        new_NoPunc = str()
        for i in NoPunc:
            new_NoPunc += i
        if Speaker != Spr or Flag == True:
            #  Làm gì đó
            Spr = Speaker
            SL = Silent
            ReDict['Speaker'].append(Speaker)
            ReDict['khong dau'].append(new_NoPunc)
            ReDict['co dau'].append(new_Punc)
            if XD == True:
                ReDict['Xuong Dong'].append(str(Silent[2]))
                Flag = False
        else: 
            ReDict['khong dau'][-1] += ' ' + new_NoPunc
            ReDict['co dau'][-1] += ' ' + new_Punc
            # commitCount+=1
            # if commitCount==newLine:
            #     commitCount=0
            #     ReDict['co dau'][-1] += '\n'
            if XD == True:
                ReDict['Xuong Dong'][-1] += ' ' + str(Silent[2])
        if XD == True:
            if Silent[2] == True:
                Flag = True
    return ReDict



def add_punc(final, final_with_punc):
    # x = recognize(final)
    x = final
    y = final_with_punc
    new = []
    for i in range(len(x)):
        if ':' in x[i]:
            t = x[i].split(':')[1].strip()
            new.append(x[i].replace(t, y[i]))
        else:
            new.append(y[i])
    return new


def recognize(conser):

    def capital(noun):
        return ' '.join([str(wordx).capitalize() for wordx in noun.split()])

    final = conser
    military_rank = ['đại tướng', 'trung tướng', 'thiếu tướng', 'đại tá','thượng tá', 'trung tá', 'thiếu tá', 'đại úy', 'thượng úy', 'trung úy', 'thiếu úy']
    start_para = [ 'xin trân trọng kính','tôi xin kính','xin kính','tôi xin', 'xin', 'kính','tôi']
    vocative = ['đồng chí', 'đại biểu']
    checkpoint_start = ['phát biểu','đặt câu hỏi','trả lời','cho ý kiến','nêu ý kiến']
    end_para = ['tôi xin báo cáo hết', 'xin phép báo cáo hết' ,'xin báo cáo hết'
               , 'xin phép hết','báo cáo hết','tôi xin hết', 'xin hết']
    thank_full = ['tôi xin cảm ơn','tôi cảm ơn','cảm ơn']

    end_co=[]
    for i in end_para:
        for j in thank_full:
            end = i+' '+j
            end_co.append(end)
    pos = pot(final)
    #nhận diện chỉ tử mời
    for i, (wor, tag) in enumerate(pos):
        if wor == 'mời' and i != len(pos)-1 and i != 0:
            if pos[i+1][0] in vocative and str(pos[i-1][0]).lower() not in start_para:
                original = pos[i-1][0] + ' mời'
                subtitute = pos[i-1][0] + ' xin mời'
                final = final.replace( original, subtitute)
            elif pos[i+1][0] in military_rank and str(pos[i-1][0]).lower() not in start_para:
                original = pos[i-1][0] + ' mời ' + pos[i+1][0]
                subtitute = pos[i-1][0] + ' xin mời đồng chí'
                final = final.replace(original, subtitute)
    words = []
    tags = []
    # Thay thế tất cả start bằng xin_mời
    for start in start_para:
        find = re.compile(start+' mời')
        for m in find.finditer(final.lower()):
            sta_in , end_in = m.start(), m.end()
            final = final[0:sta_in] + 'xin_mời' + final[end_in:]
    #Thay thế các end_co bằng báo cáo hết
    for end in end_co:
        find = re.compile(end)
        for m in find.finditer(final.lower()):
            sta_in, end_in = m.start(), m.end()
            final = final[0:sta_in] + 'báo_cáo_hết' + final[end_in:]
    #Thay thế tất cả end bằng báo_cáo_hết
    for end in end_para:
        find = re.compile(end)
        for m in find.finditer(final.lower()):
            sta_in, end_in = m.start(), m.end()
            final = final[0:sta_in] + 'báo_cáo_hết' + final[end_in:]


    # Lưu từ và phân loại vào list words, tags
    postag = pot(final)
    for text, tag in postag:
        words.append(text)
        tags.append(tag)
    # Xử lý văn bản
    state = []
    for i, word in enumerate(words):
        # Xét trường hợp không phải những từ cuối câu
        if i != len(words)-1 and i != len(words)-2:
            # xét từ đồng chí và không có rank
            if word in vocative and words[i+1].lower() not in military_rank:
                # từ liền trước là xin mời
                if words[i-1] == 'xin_mời':
                    # nếu từ liền sau là N hoặc Np
                    if tags[i+1] in ['N', 'Np', 'V'] or len(words[i+1]) > 5:
                        # print( 'MĐ: xin mời k có rank '+ words[i+1] + ' TH1' )
                        ori_sta = 'xin_mời '+word+' '+ words[i+1]
                        sub_sta = '==Đồng chí ' + capital(words[i+1]) + ': '
                        if ori_sta + sub_sta in state:
                            pass
                        else:
                            final = final.replace(ori_sta, ori_sta + sub_sta)
                            state.append(ori_sta + sub_sta)
                    # nếu từ liền sau tiếp là N hoặc Np
                    elif tags[i+2] in ['N', 'Np', 'V']  or len(words[i+2]) > 5:
                        # print( 'MĐ: xin mời k có rank '+ words[i+1] + ' TH2' )
                        ori_sta = 'xin_mời '+word+' '+ str(words[i+1]) + ' ' + words[i+2]
                        sub_sta = '==Đồng chí ' + capital(words[i+1]) + ' ' + words[i+2] +':'
                        if ori_sta + sub_sta in state:
                            pass
                        else:
                            final = final.replace(ori_sta, ori_sta + sub_sta)
                            state.append(ori_sta + sub_sta)
            # xét từ rank
            elif word in vocative and words[i+1].lower() in military_rank:
                if words[i-1] == 'xin_mời':
                    # nếu từ liền sau là N hoặc Np
                    if tags[i+2] in ['N','Np','V'] or len(words[i+2]) > 5:
                        # print( 'MĐ: xin mời '+ words[i+1] + ' TH1' )
                        ori_sta = 'xin_mời '+word+' '+ words[i+1] + ' ' + words[i+2]
                        sub_sta = '=='  + str(words[i+1]) + ' ' + capital(words[i+2]) + ': '
                        if ori_sta + sub_sta in state:
                            pass
                        else:
                            final = final.replace(ori_sta, ori_sta + sub_sta)
                            state.append(ori_sta + sub_sta)
                    # nếu từ liền sau tiếp là N hoặc Np
                    elif tags[i+3] == 'N' or tags[i+3] == 'Np'  or len(words[i+3]) > 5:
                        # print( 'MĐ: xin mời '+ words[i+1] + ' TH2' )
                        ori_sta = 'xin_mời '+word+' '+ words[i+1] +' ' + str(words[i+2]) + ' ' + words[i+3]
                        sub_sta = '==' + words[i+1] + ' ' + capital(words[i+2]) + ' ' + capital(words[i+3]) +':'
                        if ori_sta + sub_sta in state:
                            pass
                        else:
                            final = final.replace(ori_sta, ori_sta + sub_sta)
                            state.append(ori_sta + sub_sta)
            elif word in military_rank:
                if words[i-1] == 'xin_mời':
                    if tags[i+1] in ['N','Np','V'] or len(words[i+1]) > 5:
                        # print('MĐ không có từ đồng chí TH1')
                        ori_sta = 'xin_mời '+word+' '+ words[i+1]
                        sub_sta = '=='+ word + ' ' + capital(words[i+1])+': '
                        if ori_sta + sub_sta in state:
                            pass
                        else:
                            final = final.replace(ori_sta, ori_sta + sub_sta)
                            state.append(ori_sta + sub_sta)
                    # nếu từ liền sau tiếp là N hoặc Np
                    elif tags[i+2] == 'N' or tags[i+2] == 'Np'  or len(words[i+2]) > 5:
                        # print('MĐ không có từ đồng chí TH1')
                        ori_sta = 'xin_mời '+word+' '+ words[i+1] +' ' + str(words[i+2])
                        sub_sta = '==' + capital(words[i+1]) + ' ' + capital(words[i+2]) + ':'
                        if ori_sta + sub_sta in state:
                            pass
                        else:
                            final = final.replace(ori_sta, ori_sta + sub_sta)
                            state.append(ori_sta + sub_sta)
    final = final.replace('báo_cáo_hết','==')
    final = final.replace('xin_mời','xin mời')
    for check in checkpoint_start:
        find = re.compile(': '+ check)
        for m in find.finditer(final.lower()):
            sta_in, end_in = m.start(), m.end()
            final = final[0:sta_in] + ': cho_ý_kiến' + final[end_in:]
    final = final.replace(': cho_ý_kiến', ':')

    paragraph =[ para.strip() for para in final.split('==')]
    return paragraph


# # input_1 with file json
# def process_input(input_1):
#     text_no_punc = input_1['khong dau']
#     text_punc =input_1['co dau']
#     text_done = []
#     index_replace =[]
#     text_process = '/n '.join(text_no_punc)
#     #print(text_process)
#     text_reg = recognize(text_process)
#     for text in text_reg:
#         if '/n ' in text:
#             a = text.split('/n ')
#             for para in a:
#                 text_done.append(para)
#         if '/n ' not in text:
#             text_done.append(text)
#     for i in range(len(text_done)):
#         if ':' in text_done[i] and i != len(text_done)-1:
#             # print(text_done[i])
#             x = text_done[i] + ' ' + text_done[i+1]
#             index_replace.append(i+1)
#             text_done[i] = x
#     minus=0
#     for i in index_replace:
#         i -= minus
#         del text_done[i]
#         minus += 1
#     while '' in text_done:
#         text_done.remove("")
#     while '' in text_punc:
#         text_punc.remove("")
#     # print(len(text_done), len(text_punc))
#     if len(text_done) == len(text_punc):
#         final_text = add_punc(text_done, text_punc)
#         a = 'Success'
#     else:
#         a = 'Make a change'
#         for i in range(len(text_done)):
#             if text_no_punc[i].lower() not in text_done[i].lower():
#                 # print(text_no_punc[i],'\n\n\n\n' ,text_done[i])
#                 text_no_punc.insert(i,text_done[i])
#                 text_punc.insert(i,text_done[i])
#         # print(len(text_done), len(text_punc))
#         final_text = add_punc(text_done,text_punc)
#     return a, final_text
def process_input(input_1):
    text_no_punc = input_1['khong dau']
    text_done = []
    index_replace =[]
    text_process = ' ... '.join(text_no_punc)
    text_reg = recognize(text_process)
    for text in text_reg:
        text = text.replace(':  ...',':')
        if '... ' in text:
            a = text.split('... ')
            for para in a:
                if ':' in para:
                    text_done.append(para)
                elif '' == para:
                    pass
                else:
                    para = 'Vô danh: ' + para
                    text_done.append(para)
        if '... ' not in text:
            text_done.append(text)
    final_text = text_done
    return final_text


def split_silence(audio, speaker, silence_time=300,time_ok=None,match_amp=False,set_channel=False):
    def export_audio(audio, count, name):
        audios = audio.set_frame_rate(16000)
        audios.export(os.path.join(name + '/file_{}.wav'.format(str(count))), format='wav')
    def match_target_amplitude(sound,target_dBFS):
        change_in_dBFS = target_dBFS - sound.dBFS
        return sound.apply_gain(change_in_dBFS)
    folder = Thu_muc_cho_audio_cat_theo_silence
    os.makedirs(folder, exist_ok=True)
    name = audio.split('/')
    name = name[-2] + "/" + name[-1].replace('.wav', '')
    name = os.path.join(folder, name)
    os.makedirs(name, exist_ok=True)

    myaudio = AudioSegment.from_file(audio, "wav")
    if match_amp == True:
        myaudio = match_target_amplitude(myaudio,-20.0)
    if set_channel == True:
        myaudio = myaudio.set_channels(1)
    dbfs = myaudio.dBFS
    duration_in_sec = len(myaudio) / 1000

    mydict = dict()
    t_dict = dict()
    mylist = []

    # Bien = (myaudio.max_dBFS - dbfs -1.5)
    # threshold = dbfs - Bien / 2 

    # Lấy các khoảng silence trong audio
    if match_amp == True:
        silences = silence.detect_silence(myaudio,
                                        min_silence_len=silence_time,
                                        silence_thresh=dbfs-16)
    else:
        Bien = (myaudio.max_dBFS - dbfs -1.5)
        threshold = dbfs - Bien / 2 
        silences = silence.detect_silence(myaudio,
                                        min_silence_len=silence_time,
                                        silence_thresh=dbfs-16)
    silences = [((start / 1000), (stop / 1000)) for start, stop in silences]
    # print(silences)

    if len(silences) > 0:
        n_silence = []
        if silences[0][0] == 0.0:
            n_silence.append(silences[0])
            silences.pop(0)

        # Chỉnh lại, làm tròn sec
        for i in silences:
            if round(i[0]) < i[0]:
                temp = (i[0] - 0.1, i[1])
            else:
                temp = (round(i[0]), i[1])
            temp = (i[0] - 0.1, i[1])
            n_silence.append(temp)

        count = 1
        sp = 0.0
        for start, end in silences:
            temp = name + '/file_' + str(
                count) + '.wav'  # vị trí file: lưu file ở thư mục nào thì địa chỉ tới thư mục đó
            t_dict[temp] = (end, start)
            ep= start
            es= end
            ss= start
            lists = [temp, sp, ep, es, ss]
            sp= end
            mylist.append(lists)
            if time_ok != None:
                t_dict[temp] = (time_ok + start * 1000,time_ok +end * 1000)

            count += 1

        start = 0.0
        end = duration_in_sec
        s_audio = myaudio[start * 1000:n_silence[0][0] * 1000]
        export_audio(s_audio, 1, name)
        count = 2
        for i in range(len(n_silence) - 1):
            s = n_silence[i][1]
            e = n_silence[i + 1][0]
            n_audio = myaudio[s * 1000:e * 1000]
            export_audio(n_audio, count, name)
            count += 1
        if n_silence[len(n_silence) - 1][1] != end:
            e_audio = myaudio[n_silence[len(n_silence) - 1][1] * 1000:end * 1000]
            export_audio(e_audio, count, name)
            temp = name + '/file_' + str(count) + '.wav'
            t_dict[temp] = (0,0)
    else:
        temp = name + '/file_' + str(0) + '.wav'
        export_audio(myaudio, 0, name)
        mylist.append([temp,0,duration_in_sec,0,0])
        t_dict[temp] = (0,0)
    mydict[speaker] = t_dict
    return mydict, mylist


def Xuong_dong(indict,inlist):
    Adict = indict.copy()
    Fkey = list(Adict)[0]
    Adict = Adict[Fkey]
    new_dict = dict()
    for index,x in enumerate(zip(Adict.items(),inlist)):
        bdict,elelist = x
        addr,time_tuple = bdict
        time_tuple = list(time_tuple)
        time_tuple.append(elelist[-1])
        time_tuple = tuple(time_tuple)
        new_dict[addr]=time_tuple
    Rdict = {Fkey:new_dict}
    return Rdict

def Process(My_info,My_info_list,time=10):
    def checkNewLine(list,time):
        Silence = 0
        for index in list:
            if len(index) == 5:
                index.insert(5,False)
            Silence += abs(index[2]-index[1])
            if(Silence>=time):
                index[5]=True
                Silence = 0
        return list
    for i in range(len(My_info_list)):
        My_info_list[i] = checkNewLine(My_info_list[i],time)
    for i in range(len(My_info)):
        My_info[i] = Xuong_dong(My_info[i],My_info_list[i])
    return My_info
    
def downline(final_text):
    for i in range(len(final_text)):
        count = 0
        start = 0
        while 1:
            try:
                start = final_text[i].index('.',start)
                count += 1
            except:
                break
            if count == 2:
                final_text[i] = final_text[i][:start] + '\n' + final_text[i][start:]
                count = 0